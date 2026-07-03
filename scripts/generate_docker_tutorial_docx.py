"""Generate Docker tutorial as a Microsoft Word document."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_code_style(paragraph):
    for run in paragraph.runs:
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")


def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F5F5F5")
    run._element.rPr.append(shading)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
    for row_idx, row in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row):
            row_cells[col_idx].text = cell_text
    doc.add_paragraph()


def main():
    doc = Document()

    title = doc.add_heading("Dockerizing Your FastAPI + SQLAlchemy Project", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph(
        "A hands-on tutorial tailored to the fastapi-sqlalchemy Product API "
        "(FastAPI, SQLAlchemy, SQLite, and uv)."
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Table of Contents
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "1. What You Are Doing (and Why)",
        "2. Prerequisites",
        "3. How Docker Fits This Project",
        "4. Step 1 — Prepare the App for Containers",
        "5. Step 2 — Create .dockerignore",
        "6. Step 3 — Write the Dockerfile",
        "7. Step 4 — Build the Image",
        "8. Step 5 — Run the Container",
        "9. Step 6 — Add docker-compose.yml",
        "10. Step 7 — Test the Containerized API",
        "11. How It All Fits Together",
        "12. Best Practices Summary",
        "13. Common Mistakes to Avoid",
        "14. Teaching Checklist",
    ]
    for item in toc_items:
        doc.add_paragraph(item, style="List Number")
    doc.add_page_break()

    # Section 1
    doc.add_heading("1. What You Are Doing (and Why)", level=1)
    doc.add_paragraph(
        "Goal: Package your app so it runs the same way on any machine — yours, "
        "a teammate's, or a cloud server — without \"it works on my machine\" problems."
    )
    doc.add_heading("What Docker gives you", level=2)
    add_table(
        doc,
        ["Concept", "Plain English"],
        [
            ["Image", "A read-only recipe: OS + Python + your code + dependencies"],
            ["Container", "A running instance of that image (like uvicorn main:app, but isolated)"],
            ["Volume", "Persistent storage outside the container (critical for SQLite)"],
            ["Port mapping", "8000:8000 = host port 8000 → container port 8000"],
        ],
    )
    doc.add_heading("Your app today", level=2)
    add_code_block(
        doc,
        "fastapi-sqlalchemy/\n"
        "├── main.py          ← entry point (uvicorn serves main:app)\n"
        "├── database.py      ← SQLite at ./products.db (hardcoded)\n"
        "├── models.py, schemas.py, crud.py\n"
        "├── pyproject.toml   ← dependencies\n"
        "├── uv.lock          ← pinned versions (reproducible installs)\n"
        "└── products.db      ← local SQLite file",
    )
    doc.add_paragraph(
        "Docker will not replace your code. It wraps how the app is installed and started."
    )

    # Section 2
    doc.add_heading("2. Prerequisites", level=1)
    doc.add_heading("On your machine", level=2)
    doc.add_paragraph("1. Docker Desktop (Windows): https://docs.docker.com/desktop/setup/install/windows-install/", style="List Number")
    doc.add_paragraph("2. Your app already runs locally:", style="List Number")
    add_code_block(
        doc,
        "cd C:\\Users\\HP\\Documents\\fastapi-sqlalchemy\n"
        ".\\.venv\\Scripts\\Activate.ps1\n"
        "uvicorn main:app --reload",
    )
    doc.add_paragraph("3. Verify Docker:", style="List Number")
    add_code_block(doc, "docker --version\ndocker compose version")
    doc.add_paragraph("You should see version strings for both.")

    # Section 3
    doc.add_heading("3. How Docker Fits This Project", level=1)
    doc.add_paragraph("Three project-specific decisions matter:")
    doc.add_heading("A. SQLite and persistence", level=2)
    doc.add_paragraph(
        "Your database is a file (products.db). Container filesystem is ephemeral — "
        "when the container is removed, that file is gone unless you use a volume."
    )
    doc.add_paragraph("Best practice: Store the DB on a mounted volume, e.g. /app/data/products.db.")
    doc.add_heading("B. Hardcoded DATABASE_URL", level=2)
    doc.add_paragraph(
        "Right now database.py contains DATABASE_URL = \"sqlite:///./products.db\". "
        "That path works locally but is awkward in Docker. Use an environment variable "
        "so the same image works locally and in containers."
    )
    doc.add_heading("C. You use uv, not requirements.txt", level=2)
    doc.add_paragraph(
        "Your lockfile (uv.lock) pins exact versions. The Dockerfile should use "
        "uv sync --frozen so container builds match your dev environment."
    )

    # Section 4
    doc.add_heading("4. Step 1 — Prepare the App for Containers", level=1)
    doc.add_paragraph(
        "Why: Configuration (database path, ports) should come from the environment "
        "in containers, not from code baked into the image."
    )
    doc.add_heading("Update database.py", level=2)
    doc.add_paragraph("Replace the top of the file with:")
    add_code_block(
        doc,
        'import os\n\n'
        "from sqlalchemy import create_engine\n"
        "from sqlalchemy.orm import sessionmaker, declarative_base\n\n"
        'DATABASE_URL = os.getenv("DATABASE_URL", "sqlite:///./products.db")\n\n'
        "_connect_args = (\n"
        '    {"check_same_thread": False}\n'
        '    if DATABASE_URL.startswith("sqlite")\n'
        "    else {}\n"
        ")\n\n"
        "engine = create_engine(DATABASE_URL, connect_args=_connect_args)\n\n"
        "SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)\n\n"
        "Base = declarative_base()\n\n\n"
        "def get_db():\n"
        "    db = SessionLocal()\n"
        "    try:\n"
        "        yield db\n"
        "    finally:\n"
        "        db.close()",
    )
    add_table(
        doc,
        ["Change", "Why"],
        [
            ["os.getenv(\"DATABASE_URL\", ...)", "Default stays the same locally; Docker overrides via env"],
            ["Conditional connect_args", "check_same_thread is SQLite-only; needed if you switch to PostgreSQL later"],
        ],
    )
    doc.add_paragraph("Local behavior is unchanged — no env var means ./products.db as before.")
    doc.add_heading("Why not change main.py?", level=2)
    doc.add_paragraph(
        "uvicorn main:app stays the same. Docker only changes how Uvicorn is invoked "
        "(--host 0.0.0.0 inside the container)."
    )

    # Section 5
    doc.add_heading("5. Step 2 — Create .dockerignore", level=1)
    doc.add_paragraph(
        "Why: Tells Docker what not to copy into the image. Smaller builds, faster context "
        "upload, no secrets or junk."
    )
    doc.add_paragraph("Create .dockerignore in the project root:")
    add_code_block(
        doc,
        "# Virtual environment — dependencies are installed inside the image\n"
        ".venv/\n\n"
        "# Python cache\n"
        "__pycache__/\n"
        "*.py[cod]\n\n"
        "# Local SQLite — use a volume in Docker instead\n"
        "products.db\n\n"
        "# Git\n"
        ".git/\n"
        ".gitignore\n\n"
        "# IDE / OS\n"
        ".vscode/\n"
        ".idea/\n"
        "*.swp\n"
        "Thumbs.db\n\n"
        "# Docs (optional)\n"
        "TUTORIAL.md\n"
        "README.md",
    )
    doc.add_paragraph(
        "Teaching point: .dockerignore is like .gitignore, but for the Docker build context. "
        "Without it, Docker might copy your entire .venv (hundreds of MB) on every build."
    )

    # Section 6
    doc.add_heading("6. Step 3 — Write the Dockerfile", level=1)
    doc.add_paragraph("Why: The Dockerfile is the step-by-step recipe for your image.")
    doc.add_paragraph("Create Dockerfile in the project root:")
    add_code_block(
        doc,
        "# Base image: official Python 3.13 on Debian slim\n"
        "FROM python:3.13-slim\n\n"
        "# Install uv (fast, reproducible installs from uv.lock)\n"
        "COPY --from=ghcr.io/astral-sh/uv:latest /uv /uvx /bin/\n\n"
        "ENV PYTHONDONTWRITEBYTECODE=1 \\\n"
        "    PYTHONUNBUFFERED=1\n\n"
        "WORKDIR /app\n\n"
        "COPY pyproject.toml uv.lock ./\n"
        "RUN uv sync --frozen --no-dev --no-install-project\n\n"
        "COPY main.py database.py models.py schemas.py crud.py ./\n\n"
        "RUN mkdir -p /app/data\n\n"
        "ENV DATABASE_URL=sqlite:////app/data/products.db\n\n"
        "EXPOSE 8000\n\n"
        'CMD ["uv", "run", "uvicorn", "main:app", "--host", "0.0.0.0", "--port", "8000"]',
    )
    doc.add_heading("Line-by-line teaching notes", level=2)
    add_table(
        doc,
        ["Instruction", "What it does", "Why"],
        [
            ["FROM python:3.13-slim", "Base OS + Python 3.13", "Matches requires-python in pyproject.toml"],
            ["COPY --from=ghcr.io/.../uv", "Multi-stage copy of uv binary", "Official pattern from Astral"],
            ["PYTHONDONTWRITEBYTECODE=1", "Skips .pyc files", "Smaller image, fewer cache issues"],
            ["PYTHONUNBUFFERED=1", "Stdout/stderr unbuffered", "docker logs shows output immediately"],
            ["WORKDIR /app", "Default directory in container", "Paths in CMD and COPY are relative to here"],
            ["COPY pyproject.toml + RUN uv sync", "Install deps before app code", "Layer caching: code changes won't reinstall deps"],
            ["uv sync --frozen --no-dev", "Install exactly what uv.lock says", "Reproducible builds"],
            ["sqlite:////app/data/products.db", "Four slashes = absolute path", "SQLite URL format for absolute paths"],
            ["EXPOSE 8000", "Documents the port", "Does not publish to host; docker run -p does"],
            ["--host 0.0.0.0", "Listen on all interfaces", "Default 127.0.0.1 rejects outside connections"],
            ["No --reload", "Single stable process", "--reload is for local dev only"],
        ],
    )

    # Section 7
    doc.add_heading("7. Step 4 — Build the Image", level=1)
    doc.add_paragraph(
        "Why: docker build executes the Dockerfile and produces a named image you can run many times."
    )
    doc.add_paragraph("From the project root:")
    add_code_block(
        doc,
        "cd C:\\Users\\HP\\Documents\\fastapi-sqlalchemy\n"
        "docker build -t fastapi-sqlalchemy:latest .",
    )
    add_table(
        doc,
        ["Part", "Meaning"],
        [
            ["docker build", "Run the Dockerfile instructions"],
            ["-t fastapi-sqlalchemy:latest", "Tag (name:version) for the image"],
            [".", "Build context = current directory (respects .dockerignore)"],
        ],
    )
    doc.add_paragraph("Verify:")
    add_code_block(doc, "docker images fastapi-sqlalchemy")

    # Section 8
    doc.add_heading("8. Step 5 — Run the Container", level=1)
    doc.add_paragraph(
        "Why: A running container is your API in an isolated environment with its own filesystem and network."
    )
    doc.add_heading("Basic run (with persistence)", level=2)
    add_code_block(
        doc,
        "docker run --rm -p 8000:8000 `\n"
        "  -v fastapi-sqlalchemy-data:/app/data `\n"
        "  --name product-api `\n"
        "  fastapi-sqlalchemy:latest",
    )
    add_table(
        doc,
        ["Flag", "Meaning"],
        [
            ["--rm", "Remove container when it stops (good for demos)"],
            ["-p 8000:8000", "Host port 8000 → container port 8000"],
            ["-v fastapi-sqlalchemy-data:/app/data", "Named volume → SQLite survives container restarts"],
            ["--name product-api", "Friendly name for docker stop product-api"],
            ["fastapi-sqlalchemy:latest", "Image to run"],
        ],
    )
    doc.add_paragraph("Stop the container (in another terminal):")
    add_code_block(doc, "docker stop product-api")
    doc.add_paragraph("Data remains in volume fastapi-sqlalchemy-data.")

    # Section 9
    doc.add_heading("9. Step 6 — Add docker-compose.yml", level=1)
    doc.add_paragraph(
        "Why: docker run with many flags is hard to remember. Compose defines services in YAML — "
        "one command to start everything."
    )
    add_code_block(
        doc,
        "services:\n"
        "  api:\n"
        "    build: .\n"
        "    image: fastapi-sqlalchemy:latest\n"
        "    container_name: product-api\n"
        "    ports:\n"
        '      - "8000:8000"\n'
        "    environment:\n"
        "      DATABASE_URL: sqlite:////app/data/products.db\n"
        "    volumes:\n"
        "      - product-data:/app/data\n"
        "    restart: unless-stopped\n\n"
        "volumes:\n"
        "  product-data:",
    )
    doc.add_heading("Commands", level=2)
    add_code_block(
        doc,
        "# Build and start in background\n"
        "docker compose up --build -d\n\n"
        "# View logs\n"
        "docker compose logs -f api\n\n"
        "# Stop and remove containers (volume kept)\n"
        "docker compose down\n\n"
        "# Stop and remove containers AND volume (wipes DB)\n"
        "docker compose down -v",
    )

    # Section 10
    doc.add_heading("10. Step 7 — Test the Containerized API", level=1)
    doc.add_paragraph("With the container running:")
    doc.add_heading("Browser", level=2)
    doc.add_paragraph("Swagger UI: http://localhost:8000/docs", style="List Bullet")
    doc.add_paragraph("ReDoc: http://localhost:8000/redoc", style="List Bullet")
    doc.add_heading("PowerShell", level=2)
    add_code_block(
        doc,
        'Invoke-RestMethod -Uri "http://localhost:8000/products" -Method POST `\n'
        '  -ContentType "application/json" `\n'
        '  -Body \'{"name":"Docker Mug","description":"Container-themed mug","price":14.99,"stock":42}\'\n\n'
        'Invoke-RestMethod -Uri "http://localhost:8000/products" -Method GET',
    )
    doc.add_heading("Prove persistence", level=2)
    for step in [
        "Create a product.",
        "Run docker compose down (no -v).",
        "Run docker compose up -d.",
        "GET /products — data should still be there.",
    ]:
        doc.add_paragraph(step, style="List Number")

    # Section 11
    doc.add_heading("11. How It All Fits Together", level=1)
    add_code_block(
        doc,
        "Your machine (host)\n"
        "  Browser / curl  →  localhost:8000\n"
        "                         │\n"
        "                         ▼  (-p 8000:8000)\n"
        "  Docker container: product-api\n"
        "    uv run uvicorn main:app --host 0.0.0.0 --port 8000\n"
        "         │\n"
        "         ▼\n"
        "    main.py → crud.py → SQLAlchemy\n"
        "         │\n"
        "         ▼\n"
        "    /app/data/products.db  ◄── volume: product-data",
    )
    add_table(
        doc,
        ["Phase", "What happens"],
        [
            ["docker build", "Install Python deps from uv.lock; copy .py files"],
            ["docker run", "Start Uvicorn; create_all() creates tables on startup"],
            ["Volume", "SQLite file outlives container restarts"],
        ],
    )

    # Section 12
    doc.add_heading("12. Best Practices Summary", level=1)
    add_table(
        doc,
        ["Practice", "Your project"],
        [
            ["Use env vars for config", "DATABASE_URL in database.py"],
            [".dockerignore", "Exclude .venv, products.db, __pycache__"],
            ["Pin dependencies", "uv.lock + uv sync --frozen"],
            ["Layer caching", "Copy pyproject.toml / uv.lock before app code"],
            ["Slim base image", "python:3.13-slim"],
            ["No --reload in image", "Production CMD without reload"],
            ["Bind 0.0.0.0 in container", "Required for port mapping to work"],
            ["Persist SQLite", "Named volume on /app/data"],
            ["Compose for teams", "One file, one docker compose up"],
        ],
    )

    # Section 13
    doc.add_heading("13. Common Mistakes to Avoid", level=1)
    add_table(
        doc,
        ["Mistake", "Symptom", "Fix"],
        [
            ["Forgot --host 0.0.0.0", "Connection refused from host", "Add to CMD"],
            ["No volume for SQLite", "Data disappears after docker rm", "Use -v or Compose volumes"],
            ["Copied .venv into image", "Huge image, slow builds", "Add .dockerignore"],
            ["Used --reload in production", "Extra processes, odd restarts", "Remove from CMD"],
            ["Wrong SQLite URL", "DB not where you expect", "Use sqlite:////absolute/path (4 slashes)"],
            ["Only EXPOSE, no -p", "Can't reach API from host", 'ports: "8000:8000"'],
            ["Edited code, old behavior", "Stale image", "docker compose up --build"],
        ],
    )

    # Section 14
    doc.add_heading("14. Teaching Checklist", level=1)
    checklist = [
        "Explain image vs container vs volume (Section 1)",
        "Confirm app runs locally with uvicorn",
        "Update database.py for DATABASE_URL env var",
        "Add .dockerignore and explain why",
        "Write Dockerfile line by line",
        "docker build -t fastapi-sqlalchemy:latest .",
        "docker run with -p and -v; hit /docs",
        "Create a product; restart container; verify data persists",
        "Introduce docker-compose.yml and docker compose up --build -d",
        "Discuss production next steps (Postgres, health checks, CI)",
    ]
    for item in checklist:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run("☐ ").bold = True
        p.add_run(item)

    doc.add_heading("Final layout after Dockerizing", level=2)
    add_code_block(
        doc,
        "fastapi-sqlalchemy/\n"
        "├── Dockerfile\n"
        "├── docker-compose.yml\n"
        "├── .dockerignore\n"
        "├── database.py          ← reads DATABASE_URL from env\n"
        "├── main.py\n"
        "├── models.py, schemas.py, crud.py\n"
        "├── pyproject.toml\n"
        "├── uv.lock\n"
        "└── products.db          ← local dev only; Docker uses volume",
    )

    doc.add_heading("Quick command reference", level=2)
    add_code_block(
        doc,
        "# Build\n"
        "docker build -t fastapi-sqlalchemy:latest .\n\n"
        "# Run (manual)\n"
        "docker run --rm -p 8000:8000 -v fastapi-sqlalchemy-data:/app/data fastapi-sqlalchemy:latest\n\n"
        "# Compose (recommended day-to-day)\n"
        "docker compose up --build -d\n"
        "docker compose logs -f api\n"
        "docker compose down",
    )

    output_path = "DOCKER_TUTORIAL.docx"
    doc.save(output_path)
    print(f"Created: {output_path}")


if __name__ == "__main__":
    main()
